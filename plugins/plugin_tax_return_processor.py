"""
MC & S Plugin: Tax Return Deductions Processor
================================================
Plugin ID  : plugin_tax_return_processor
Version    : 1.0.0

WHAT IT DOES
------------
Monitors the inbox for DOCUMENTS_RECEIVED emails (flagged by the Email Triage
plugin). For each flagged email, it:

  1. Downloads all attachments (PDFs, images, spreadsheets)
  2. Extracts deduction items using Claude AI
  3. Analyses each item against ATO rules (D1-D10)
  4. Generates a Word (.docx) deductions report with:
       - Allowed items table with amounts and work %
       - Disallowed items table with ATO CITATION for each rejection
         (Fix 3: Citation-Linked Reasoning — no more "black box" decisions)
       - Items flagged for Elio's professional judgment
       - Suggested client follow-up questions
  5. Emails the report to elio@mcands.com.au

CITATION REQUIREMENT (Fix 3 from EVA friction analysis)
---------------------------------------------------------
Every disallowed item MUST include a specific citation from:
  - ITAA97 section (e.g., "s8-1(2)(b) — private or domestic in nature")
  - ATO Taxation Ruling (e.g., "TR 2022/1, para 45")
  - ATO Practical Compliance Guideline (e.g., "PCG 2023/1, cl 14")
  - ATO Individual Tax Return Instructions 2025 (e.g., "Instructions 2025, D3, p.47")

This transforms EVA from a black box into a transparent, verifiable research
assistant that accountants can trust and audit.

SCHEDULE
--------
Runs every 5 minutes to check for newly flagged emails.
"""

import json
import os
import re
import traceback
from datetime import datetime
from io import BytesIO

import anthropic

from plugin_base import AgentPlugin, PluginContext, PluginResult, Schedule


# ── Prompt templates ──────────────────────────────────────────────────────────

EXTRACTION_SYSTEM_PROMPT = """\
You are a tax deductions extraction assistant for MC & S Accountants, Melbourne.

Your task is to extract ALL expense/deduction items from the provided client documents.
For each item, extract:
- description: what the expense is
- amount: dollar amount (AUD)
- category_hint: which D-category it likely belongs to (D1-D10 or "unknown")
- notes: any relevant context (dates, percentages, purpose mentioned)

Return a JSON array of objects with keys: description, amount, category_hint, notes.
If you cannot determine an amount, use null.
Be thorough — extract every expense item mentioned, even if you are unsure of deductibility.
"""

ANALYSIS_SYSTEM_PROMPT = """\
You are a senior Australian tax accountant at MC & S Accountants, Melbourne.
You are analysing deduction items for an individual tax return (FY2025).

AUTHORITY HIERARCHY (apply in this order):
1. ITAA97 — Income Tax Assessment Act 1997 (primary authority)
2. ATO Taxation Rulings (TR) and Practical Compliance Guidelines (PCG)
3. ATO Individual Tax Return Instructions 2025
4. ATO website guidance

CRITICAL CITATION REQUIREMENT:
For EVERY item you disallow, you MUST provide a specific citation. Examples:
- "s8-1(2)(b) ITAA97 — private or domestic in nature"
- "ATO Instructions 2025, D3, p.47 — conventional clothing not deductible"
- "PCG 2023/1, cl 14 — fixed rate method covers internet; cannot claim separately"
- "TR 2022/1, para 45 — effective life of laptop is 4 years"
- "s26-5 ITAA97 — penalties and fines not deductible"

Do NOT use vague reasons like "not work-related" without citing the authority.

RESPONSE FORMAT:
Return a JSON object with these keys:
{
  "allowed": [
    {
      "description": "...",
      "d_category": "D1",
      "amount": 1234.56,
      "work_pct": 80,
      "deductible_amount": 987.65,
      "notes": "..."
    }
  ],
  "disallowed": [
    {
      "description": "...",
      "amount": 123.45,
      "reason": "Conventional clothing — not occupation-specific or protective",
      "citation": "ATO Instructions 2025, D3, p.47 — conventional clothing (suits, business wear) is not deductible even if employer requires it"
    }
  ],
  "flagged_for_review": [
    {
      "description": "...",
      "amount": 456.78,
      "question": "What was the exact purchase date? Is this used exclusively for work?"
    }
  ],
  "follow_up_questions": [
    "What is the total number of work-related kilometres driven in FY2025?",
    "Please provide the exact purchase date for the laptop ($1,200)."
  ],
  "d_totals": {
    "D1": 0, "D2": 0, "D3": 0, "D4": 0, "D5": 0,
    "D6": 0, "D7": 0, "D8": 0, "D9": 0, "D10": 0
  }
}
"""

ANALYSIS_USER_TEMPLATE = """\
CLIENT: {client_name}
OCCUPATION: {occupation}
FINANCIAL YEAR: FY2025 (1 July 2024 – 30 June 2025)

EXTRACTED EXPENSE ITEMS:
{items_json}

PRIOR YEAR CONTEXT:
{prior_year_context}

Analyse each item and classify as allowed, disallowed, or flagged for review.
For every disallowed item, provide the specific ATO citation.
"""


class TaxReturnProcessorPlugin(AgentPlugin):
    """
    Processes tax return deduction documents and generates a cited report.
    """

    name        = "Tax Return Deductions Processor"
    description = "Analyses client tax documents and generates a D1-D10 deductions report with ATO citations."
    detail      = (
        "Monitors for flagged DOCUMENTS_RECEIVED emails. Downloads attachments, "
        "extracts expense items with Claude, analyses deductibility against ATO rules "
        "(ITAA97, TR, PCG, Instructions 2025), and produces a Word report. Every "
        "disallowed item includes the specific ATO citation — no more black-box decisions."
    )
    version = "1.0.0"
    icon    = "📋"

    requires_graph  = True
    requires_claude = True

    default_schedule = Schedule.every_minutes(5)

    # Track processed email IDs this session
    _processed_ids: set

    def load(self, context: PluginContext) -> bool:
        self._processed_ids = set()
        if not context.graph:
            context.log("📋 Tax Return Processor: Microsoft 365 not connected.")
            return False
        if not context.claude:
            context.log("📋 Tax Return Processor: Claude AI not configured.")
            return False
        return True

    @classmethod
    def settings_schema(cls) -> list[dict]:
        return [
            {
                "key": "report_recipient",
                "label": "Report Recipient Email",
                "default": "elio@mcands.com.au",
                "type": "text",
                "help": "Email address to send the completed deductions report to.",
            },
            {
                "key": "monitor_folder",
                "label": "Folder to Monitor",
                "default": "Inbox",
                "type": "text",
                "help": "Outlook folder to check for flagged DOCUMENTS_RECEIVED emails.",
            },
            {
                "key": "max_per_run",
                "label": "Max Emails Per Run",
                "default": "5",
                "type": "number",
                "help": "Maximum number of tax return emails to process per run.",
            },
        ]

    def run(self, context: PluginContext) -> PluginResult:
        try:
            return self._do_run(context)
        except Exception as e:
            context.log(f"📋 Tax Return Processor: Error — {e}")
            traceback.print_exc()
            return PluginResult(success=False, error=str(e))

    def _do_run(self, context: PluginContext) -> PluginResult:
        graph      = context.graph
        log        = context.log
        folder     = self.get_plugin_setting("monitor_folder", "Inbox")
        recipient  = self.get_plugin_setting("report_recipient", "elio@mcands.com.au")
        max_count  = int(self.get_plugin_setting("max_per_run", "5"))

        # Fetch flagged, unread emails (flagged by Email Triage for DOCUMENTS_RECEIVED)
        try:
            params = {
                "$filter": "isRead eq false and flag/flagStatus eq 'flagged'",
                "$top": max_count,
                "$orderby": "receivedDateTime asc",
                "$select": "id,subject,from,receivedDateTime,body,bodyPreview,hasAttachments,categories",
            }
            import requests as req
            url = f"https://graph.microsoft.com/v1.0/me/mailFolders/{folder}/messages"
            r = req.get(url, headers=graph._headers(), params=params)
            r.raise_for_status()
            emails = r.json().get("value", [])
        except Exception as e:
            return PluginResult(success=False, error=f"Could not fetch emails: {e}")

        # Filter to only DOCUMENTS_RECEIVED category (tagged by Email Triage)
        tax_emails = [
            e for e in emails
            if "Documents Received" in e.get("categories", [])
            and e["id"] not in self._processed_ids
        ]

        if not tax_emails:
            log("📋 Tax Return Processor: No flagged DOCUMENTS_RECEIVED emails to process.")
            return PluginResult(success=True, summary="No tax return emails to process.")

        log(f"📋 Tax Return Processor: Found {len(tax_emails)} email(s) to process.")
        result = PluginResult(success=True)

        for email in tax_emails:
            msg_id  = email["id"]
            subject = email.get("subject", "(No Subject)")
            sender  = email.get("from", {}).get("emailAddress", {})
            from_name  = sender.get("name", "Client")
            from_email = sender.get("address", "")

            log(f"  Processing: \"{subject}\" from {from_name}")

            try:
                # ── 1. Download attachments ───────────────────────────────────
                attachment_texts = self._get_attachment_texts(graph, msg_id, log)
                body_text = re.sub(r"<[^>]+>", " ", email.get("body", {}).get("content", ""))
                body_text = re.sub(r"\s+", " ", body_text).strip()

                all_text = body_text + "\n\n" + "\n\n".join(attachment_texts)

                if len(all_text.strip()) < 50:
                    log(f"  ⚠ No usable content found in email/attachments — skipping.")
                    self._processed_ids.add(msg_id)
                    result.items_skipped += 1
                    continue

                # ── 2. Extract expense items ──────────────────────────────────
                log(f"  Extracting expense items...")
                items = self._extract_items(context.claude, all_text[:6000])

                if not items:
                    log(f"  ⚠ No expense items found — skipping.")
                    self._processed_ids.add(msg_id)
                    result.items_skipped += 1
                    continue

                log(f"  Found {len(items)} expense item(s). Analysing deductibility...")

                # ── 3. Analyse deductibility with citations ───────────────────
                client_name = from_name
                occupation  = self._guess_occupation(body_text, items)
                analysis    = self._analyse_items(
                    context.claude, client_name, occupation, items
                )

                # ── 4. Generate Word report ───────────────────────────────────
                log(f"  Generating deductions report...")
                report_bytes, report_filename = self._generate_report(
                    client_name, occupation, subject, analysis
                )

                # ── 5. Email report to accountant ─────────────────────────────
                allowed_total = sum(
                    i.get("deductible_amount", 0) or 0
                    for i in analysis.get("allowed", [])
                )
                disallowed_count = len(analysis.get("disallowed", []))
                flagged_count    = len(analysis.get("flagged_for_review", []))

                report_subject = (
                    f"Tax Return Deductions Report — {client_name} — FY2025"
                )
                report_body = f"""
<p>Hi Elio,</p>
<p>The Tax Return Deductions Processor has completed the analysis for
<strong>{client_name}</strong> (FY2025).</p>
<table style="border-collapse:collapse;font-family:Arial;font-size:13px">
  <tr><td style="padding:4px 12px 4px 0"><strong>Total Deductible Amount:</strong></td>
      <td style="padding:4px 0"><strong>${allowed_total:,.2f}</strong></td></tr>
  <tr><td style="padding:4px 12px 4px 0">Items Allowed:</td>
      <td style="padding:4px 0">{len(analysis.get('allowed', []))}</td></tr>
  <tr><td style="padding:4px 12px 4px 0">Items Disallowed (with ATO citations):</td>
      <td style="padding:4px 0">{disallowed_count}</td></tr>
  <tr><td style="padding:4px 12px 4px 0">Items Flagged for Your Review:</td>
      <td style="padding:4px 0">{flagged_count}</td></tr>
</table>
<p>The full report is attached. Each disallowed item includes the specific
ATO citation (ITAA97 section, TR/PCG reference, or Instructions 2025 page).</p>
<p>Kind regards,<br>EVA — MC&S CoWorker</p>
"""
                # Send with attachment
                self._send_report_email(
                    graph, recipient, report_subject, report_body,
                    report_bytes, report_filename
                )

                log(f"  ✅ Report sent to {recipient} — ${allowed_total:,.2f} deductible, "
                    f"{disallowed_count} disallowed (cited), {flagged_count} flagged.")

                # Mark email as processed
                graph.mark_as_read(msg_id)
                self._processed_ids.add(msg_id)
                result.actions_taken += 1
                result.drafts_created += 1  # Counts as a report created

                self.log_activity(
                    source=self.name,
                    subject=f"Tax report: {client_name} FY2025",
                    category="tax_return",
                    action=f"Report generated — ${allowed_total:,.2f} deductible",
                    draft_created=1,
                )

            except Exception as e:
                log(f"  ❌ Error processing {from_name}: {e}")
                traceback.print_exc()
                result.items_skipped += 1

        result.summary = (
            f"{result.actions_taken} report(s) generated, "
            f"{result.items_skipped} skipped."
        )
        return result

    # ── Private helpers ───────────────────────────────────────────────────────

    def _get_attachment_texts(self, graph, msg_id: str, log) -> list[str]:
        """Download and extract text from email attachments."""
        texts = []
        try:
            import requests as req
            url = f"https://graph.microsoft.com/v1.0/me/messages/{msg_id}/attachments"
            r = req.get(url, headers=graph._headers())
            r.raise_for_status()
            attachments = r.json().get("value", [])

            for att in attachments:
                name = att.get("name", "").lower()
                content_bytes = att.get("contentBytes", "")
                if not content_bytes:
                    continue

                import base64
                raw = base64.b64decode(content_bytes)

                if name.endswith(".pdf"):
                    texts.append(self._extract_pdf_text(raw, name))
                elif name.endswith((".txt", ".csv")):
                    texts.append(raw.decode("utf-8", errors="ignore"))
                elif name.endswith((".xlsx", ".xls")):
                    texts.append(self._extract_excel_text(raw, name))
                else:
                    log(f"    Skipping unsupported attachment: {name}")

        except Exception as e:
            log(f"    ⚠ Could not fetch attachments: {e}")

        return texts

    def _extract_pdf_text(self, raw: bytes, name: str) -> str:
        """Extract text from a PDF using pdfminer or fallback."""
        try:
            import subprocess, tempfile
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                f.write(raw)
                tmp_path = f.name
            result = subprocess.run(
                ["pdftotext", tmp_path, "-"],
                capture_output=True, text=True, timeout=30
            )
            os.unlink(tmp_path)
            return f"[PDF: {name}]\n{result.stdout}"
        except Exception:
            return f"[PDF: {name} — could not extract text]"

    def _extract_excel_text(self, raw: bytes, name: str) -> str:
        """Extract text from an Excel file."""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(BytesIO(raw), data_only=True)
            lines = [f"[Excel: {name}]"]
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows(values_only=True):
                    row_text = "\t".join(str(c) for c in row if c is not None)
                    if row_text.strip():
                        lines.append(row_text)
            return "\n".join(lines)
        except Exception:
            return f"[Excel: {name} — could not extract text]"

    def _extract_items(self, claude: anthropic.Anthropic, text: str) -> list[dict]:
        """Use Claude to extract expense items from document text."""
        response = claude.messages.create(
            model=self.get_claude_model(),
            max_tokens=2000,
            system=EXTRACTION_SYSTEM_PROMPT,
            messages=[{
                "role": "user",
                "content": f"Extract all expense/deduction items from these documents:\n\n{text}"
            }],
        )
        raw = response.content[0].text.strip()
        # Strip markdown code fences if present
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        return json.loads(raw)

    def _analyse_items(
        self,
        claude: anthropic.Anthropic,
        client_name: str,
        occupation: str,
        items: list[dict],
    ) -> dict:
        """Analyse deductibility of each item with mandatory ATO citations."""
        prompt = ANALYSIS_USER_TEMPLATE.format(
            client_name=client_name,
            occupation=occupation or "Not specified",
            items_json=json.dumps(items, indent=2),
            prior_year_context="No prior year data available for this run.",
        )
        response = claude.messages.create(
            model=self.get_claude_model(),
            max_tokens=4000,
            system=ANALYSIS_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = response.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        return json.loads(raw)

    def _guess_occupation(self, body_text: str, items: list[dict]) -> str:
        """Attempt to extract occupation from document text."""
        occ_patterns = [
            r"occupation[:\s]+([A-Za-z\s]+)",
            r"employed as[:\s]+([A-Za-z\s]+)",
            r"work(?:s|ing) as[:\s]+([A-Za-z\s]+)",
        ]
        for pat in occ_patterns:
            m = re.search(pat, body_text, re.IGNORECASE)
            if m:
                return m.group(1).strip()[:50]
        return "Not specified"

    def _generate_report(
        self,
        client_name: str,
        occupation: str,
        email_subject: str,
        analysis: dict,
    ) -> tuple[bytes, str]:
        """Generate a Word (.docx) deductions report with ATO citations."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            # Fallback: generate a plain text report
            return self._generate_text_report(client_name, analysis)

        doc = Document()

        # ── Header ────────────────────────────────────────────────────────────
        title = doc.add_heading(f"Tax Return Deductions Report — FY2025", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Client: {client_name}")
        doc.add_paragraph(f"Occupation: {occupation}")
        doc.add_paragraph(f"Date Prepared: {datetime.now().strftime('%d %B %Y')}")
        doc.add_paragraph(f"Source: {email_subject}")
        doc.add_paragraph()

        # ── Allowed Items ─────────────────────────────────────────────────────
        doc.add_heading("Allowed Deductions", 1)
        allowed = analysis.get("allowed", [])
        if allowed:
            tbl = doc.add_table(rows=1, cols=5)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            for i, h in enumerate(["Description", "Category", "Amount", "Work %", "Deductible"]):
                hdr[i].text = h
                hdr[i].paragraphs[0].runs[0].bold = True

            for item in allowed:
                row = tbl.add_row().cells
                row[0].text = item.get("description", "")
                row[1].text = item.get("d_category", "")
                row[2].text = f"${item.get('amount', 0) or 0:,.2f}"
                row[3].text = f"{item.get('work_pct', 100)}%"
                row[4].text = f"${item.get('deductible_amount', 0) or 0:,.2f}"
        else:
            doc.add_paragraph("No items allowed.")

        doc.add_paragraph()

        # ── D1-D10 Totals ─────────────────────────────────────────────────────
        doc.add_heading("Deductions Summary (D1-D10)", 1)
        totals = analysis.get("d_totals", {})
        tbl2 = doc.add_table(rows=1, cols=2)
        tbl2.style = "Table Grid"
        hdr2 = tbl2.rows[0].cells
        hdr2[0].text = "Category"
        hdr2[0].paragraphs[0].runs[0].bold = True
        hdr2[1].text = "Total Deductible"
        hdr2[1].paragraphs[0].runs[0].bold = True
        for cat, amt in totals.items():
            if amt:
                row = tbl2.add_row().cells
                row[0].text = cat
                row[1].text = f"${amt:,.2f}"
        grand_total = sum(v for v in totals.values() if v)
        row = tbl2.add_row().cells
        row[0].text = "TOTAL"
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = f"${grand_total:,.2f}"
        row[1].paragraphs[0].runs[0].bold = True
        doc.add_paragraph()

        # ── Disallowed Items (with ATO citations) ─────────────────────────────
        doc.add_heading("Disallowed Items — ATO Citations", 1)
        disallowed = analysis.get("disallowed", [])
        if disallowed:
            tbl3 = doc.add_table(rows=1, cols=4)
            tbl3.style = "Table Grid"
            hdr3 = tbl3.rows[0].cells
            for i, h in enumerate(["Description", "Amount", "Reason", "ATO Citation"]):
                hdr3[i].text = h
                hdr3[i].paragraphs[0].runs[0].bold = True

            for item in disallowed:
                row = tbl3.add_row().cells
                row[0].text = item.get("description", "")
                row[1].text = f"${item.get('amount', 0) or 0:,.2f}"
                row[2].text = item.get("reason", "")
                # Citation is the key fix — displayed prominently
                citation = item.get("citation", "No citation provided")
                row[3].text = citation
        else:
            doc.add_paragraph("No items disallowed.")

        doc.add_paragraph()

        # ── Flagged for Review ────────────────────────────────────────────────
        doc.add_heading("Items Flagged for Elio's Review", 1)
        flagged = analysis.get("flagged_for_review", [])
        if flagged:
            for item in flagged:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(item.get("description", "")).bold = True
                p.add_run(f" — ${item.get('amount', 0) or 0:,.2f}")
                doc.add_paragraph(f"  Question: {item.get('question', '')}")
        else:
            doc.add_paragraph("No items flagged.")

        doc.add_paragraph()

        # ── Client Follow-Up Questions ────────────────────────────────────────
        doc.add_heading("Suggested Client Follow-Up Questions", 1)
        questions = analysis.get("follow_up_questions", [])
        for i, q in enumerate(questions, 1):
            doc.add_paragraph(f"{i}. {q}")

        # ── Save to bytes ─────────────────────────────────────────────────────
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        safe_name = re.sub(r"[^\w\s-]", "", client_name).strip().replace(" ", "_")
        filename  = f"TaxReturn_{safe_name}_FY2025_{datetime.now().strftime('%Y%m%d')}.docx"

        return buf.read(), filename

    def _generate_text_report(self, client_name: str, analysis: dict) -> tuple[bytes, str]:
        """Fallback: generate a plain text report if python-docx is unavailable."""
        lines = [
            f"TAX RETURN DEDUCTIONS REPORT — FY2025",
            f"Client: {client_name}",
            f"Date: {datetime.now().strftime('%d %B %Y')}",
            "",
            "=== ALLOWED DEDUCTIONS ===",
        ]
        for item in analysis.get("allowed", []):
            lines.append(
                f"  {item.get('d_category','?')} | {item.get('description','')} | "
                f"${item.get('deductible_amount',0):,.2f}"
            )

        lines += ["", "=== DISALLOWED ITEMS (ATO CITATIONS) ==="]
        for item in analysis.get("disallowed", []):
            lines.append(f"  {item.get('description','')} | ${item.get('amount',0):,.2f}")
            lines.append(f"    Reason: {item.get('reason','')}")
            lines.append(f"    Citation: {item.get('citation','No citation')}")

        lines += ["", "=== FLAGGED FOR REVIEW ==="]
        for item in analysis.get("flagged_for_review", []):
            lines.append(f"  {item.get('description','')} — {item.get('question','')}")

        lines += ["", "=== FOLLOW-UP QUESTIONS ==="]
        for i, q in enumerate(analysis.get("follow_up_questions", []), 1):
            lines.append(f"  {i}. {q}")

        content = "\n".join(lines).encode("utf-8")
        safe_name = re.sub(r"[^\w\s-]", "", client_name).strip().replace(" ", "_")
        filename  = f"TaxReturn_{safe_name}_FY2025_{datetime.now().strftime('%Y%m%d')}.txt"
        return content, filename

    def _send_report_email(
        self,
        graph,
        to_address: str,
        subject: str,
        body_html: str,
        attachment_bytes: bytes,
        attachment_name: str,
    ):
        """Send the report email with the Word doc attached."""
        import base64, requests as req
        encoded = base64.b64encode(attachment_bytes).decode("utf-8")

        payload = {
            "message": {
                "subject": subject,
                "body": {"contentType": "HTML", "content": body_html},
                "toRecipients": [{"emailAddress": {"address": to_address}}],
                "attachments": [
                    {
                        "@odata.type": "#microsoft.graph.fileAttachment",
                        "name": attachment_name,
                        "contentBytes": encoded,
                        "contentType": (
                            "application/vnd.openxmlformats-officedocument"
                            ".wordprocessingml.document"
                        ),
                    }
                ],
            },
            "saveToSentItems": True,
        }
        r = req.post(
            "https://graph.microsoft.com/v1.0/me/sendMail",
            headers=graph._headers(),
            json=payload,
        )
        r.raise_for_status()
