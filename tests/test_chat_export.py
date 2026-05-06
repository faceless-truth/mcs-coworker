"""
Tests for chat_export.build_chat_docx (Fix B1).

Pins the contract from task_chat_persistence_and_archive.md:
  - Build a doc from 3 fake messages, confirm bytes are non-empty.
  - Confirm Document(buf) round-trips and the heading text matches the
    input title.
"""
import io
import sys
import unittest
from datetime import datetime
from pathlib import Path

REPO_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(REPO_ROOT))

from docx import Document  # noqa: E402

from chat_export import build_chat_docx, _role_label  # noqa: E402


class BuildChatDocxTests(unittest.TestCase):
    def _three_messages(self):
        return [
            {
                "role": "user",
                "content": "What is the GST treatment of a going concern sale?",
                "created_at": "2026-05-06T14:32:00+00:00",
            },
            {
                "role": "assistant",
                "content": "A going concern sale can be GST-free if conditions are met.",
                "created_at": "2026-05-06T14:32:30+00:00",
            },
            {
                "role": "user",
                "content": "What conditions specifically?",
                "created_at": "2026-05-06T14:33:00+00:00",
            },
        ]

    def test_returns_non_empty_bytesio(self):
        buf = build_chat_docx(
            specialist_display_name="GST Specialist",
            title="GST Going Concern Sale",
            messages=self._three_messages(),
            user_display_name="Elio Scarton",
            archived_at=datetime(2026, 5, 6, 14, 35),
        )
        self.assertIsInstance(buf, io.BytesIO)
        data = buf.getvalue()
        self.assertGreater(len(data), 0)

    def test_round_trips_through_python_docx_with_matching_heading(self):
        title = "GST Going Concern Sale"
        buf = build_chat_docx(
            specialist_display_name="GST Specialist",
            title=title,
            messages=self._three_messages(),
            user_display_name="Elio Scarton",
            archived_at=datetime(2026, 5, 6, 14, 35),
        )

        doc = Document(buf)

        headings = [
            p.text
            for p in doc.paragraphs
            if p.style.name.startswith("Heading")
        ]
        self.assertIn(title, headings)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Specialist: GST Specialist", full_text)
        self.assertIn("User: Elio Scarton", full_text)
        self.assertIn("Archived: 2026-05-06 14:35", full_text)
        self.assertIn("What is the GST treatment of a going concern sale?", full_text)
        self.assertIn("A going concern sale can be GST-free", full_text)

    def test_role_labels(self):
        self.assertEqual(_role_label("user", "Elio Scarton"), "Elio Scarton:")
        self.assertEqual(_role_label("assistant", "Elio Scarton"), "MCS CoWorker:")
        self.assertEqual(_role_label("system", "Elio Scarton"), "system:")


if __name__ == "__main__":
    unittest.main()
