"""
MCS CoWorker — Flask API Server
Bridges the React/pywebview frontend to all Python backend modules.
Runs on localhost:7842 — not exposed externally.
"""
from __future__ import annotations

import html
import json
import logging
import os
import re
import sys
import threading
import traceback
from pathlib import Path

logger = logging.getLogger(__name__)
from datetime import datetime
from functools import wraps
from typing import Any

import base64
import mimetypes
import queue
import time
import uuid
from collections import deque
from flask import Flask, Response, jsonify, request, stream_with_context, send_from_directory, send_file
from flask_cors import CORS

# ── Backend imports ────────────────────────────────────────────────────────────
import config
from config import (
    init_db,
    get_setting, set_setting,
    get_all_settings,
    get_recent_activity,
    get_all_plugin_states,
    get_claude_model_fast,
    get_claude_model_reasoning,
    update_claude_models,
    get_rules, save_rule, delete_rule,
    get_staff, save_staff, delete_staff,
    get_links, save_link, delete_link,
    get_active_lessons, add_lesson, delete_lesson, toggle_lesson,
    get_style_preferences, save_style_preferences,
    add_feedback_message, get_feedback_history, clear_feedback_history,
    get_knowledge_entries, add_knowledge_entry,
    update_knowledge_entry, delete_knowledge_entry,
    get_bas_clients, get_bas_client, add_bas_client,
    update_bas_client, delete_bas_client, bulk_replace_bas_clients,
    BAS_CLIENT_COLUMNS,
)
from plugin_loader import PluginLoader
from approval_queue import ApprovalQueue
from token_meter import get_usage_summary
from event_bus import EventBus
from kpi_monitor import KPIMonitor
from specialists.registry import get_all_agents, get_agent

# ── App setup ──────────────────────────────────────────────────────────────────
app = Flask(__name__)

# Origins permitted to call the API. Used for both the general flask-cors
# config and the SSE stream's per-response Access-Control-Allow-Origin echo.
# Keep the SSE stream tight — it carries live client names/subjects, which we
# don't want any random site the accountant visits to be able to subscribe to.
ALLOWED_ORIGINS = [
    "http://localhost:7842", "http://127.0.0.1:7842",
    "http://localhost:3000", "http://127.0.0.1:3000",
    "http://localhost:5173", "http://127.0.0.1:5173",
]
CORS(app, origins=ALLOWED_ORIGINS, supports_credentials=True)
# Electron uses file:// or app:// — handle via wildcard on those routes separately

API_PORT = 7842

# ── /api/chat rate limiting ────────────────────────────────────────────────────
# Even with Fix 2 auth, the webview itself is trusted to forward requests, so a
# buggy or compromised frontend could still drain the Anthropic budget. These
# bounds put a hard ceiling on cost exposure per minute and per call.
_chat_timestamps: deque = deque()
CHAT_RATE_LIMIT = 30       # max requests
CHAT_RATE_WINDOW = 60      # per N seconds
CHAT_MAX_MESSAGE_LEN = 50000  # characters per user message
CHAT_MAX_HISTORY = 50      # conversation turns retained


def _check_chat_rate() -> bool:
    now = time.time()
    while _chat_timestamps and _chat_timestamps[0] < now - CHAT_RATE_WINDOW:
        _chat_timestamps.popleft()
    if len(_chat_timestamps) >= CHAT_RATE_LIMIT:
        return False
    _chat_timestamps.append(now)
    return True


# ── /api/chat file uploads ─────────────────────────────────────────────────────
# Specialists (GST, SMSF, Net Wealth, etc.) accept source documents — PDFs,
# spreadsheets, Word docs, images — alongside the conversation. Each upload is
# saved under DATA_DIR / "chat_uploads" keyed by a random id and referenced
# from /api/chat via {"files": [{"id", "name", "type", ...}]}.
CHAT_UPLOAD_DIR = config.DATA_DIR / "chat_uploads"
CHAT_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

MAX_UPLOAD_SIZE = 25 * 1024 * 1024   # 25 MB per file
MAX_UPLOADS_PER_MESSAGE = 5
ALLOWED_UPLOAD_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx", ".csv", ".txt",
    ".jpg", ".jpeg", ".png", ".gif",
}


def _cleanup_old_uploads():
    """Remove chat uploads older than 24 hours."""
    try:
        cutoff = time.time() - 86400
        for f in CHAT_UPLOAD_DIR.iterdir():
            try:
                if f.is_file() and f.stat().st_mtime < cutoff:
                    f.unlink(missing_ok=True)
            except Exception:
                pass
    except Exception as e:
        logger.warning(f"chat_uploads cleanup skipped: {e}")


def _schedule_upload_cleanup():
    """Run upload cleanup now and every 6 hours thereafter.

    The app is designed to run continuously, so a one-shot cleanup at
    startup is not enough — old files would otherwise accumulate.
    """
    _cleanup_old_uploads()
    timer = threading.Timer(6 * 3600, _schedule_upload_cleanup)
    timer.daemon = True
    timer.start()


_schedule_upload_cleanup()

# Shared state — populated by main.py on startup
_loader: PluginLoader | None = None
_approval_queue: ApprovalQueue | None = None
_kpi_monitor: KPIMonitor | None = None
_graph_client = None  # set by main.py after GraphClient is created
_start_time: datetime = datetime.now()


def set_loader(loader: PluginLoader):
    global _loader
    _loader = loader


def set_approval_queue(aq: ApprovalQueue):
    global _approval_queue
    _approval_queue = aq


def set_kpi_monitor(km: KPIMonitor):
    global _kpi_monitor
    _kpi_monitor = km


def set_graph_client(gc):
    global _graph_client
    _graph_client = gc


# ── API token auth ─────────────────────────────────────────────────────────────
# All /api/* routes require a per-install token. The webview injects the token
# into window.__API_TOKEN__ on load so the React frontend can send it as a
# Bearer header (or ?token= query param for EventSource, which cannot set
# custom headers). Static assets served from "/" are not gated — they have no
# side effects and gating them would break initial page load.
AUTH_EXEMPT_API_PATHS = {
    "/api/health",  # liveness probe — never sensitive
}
AUTH_EXEMPT_PREFIXES = (
    "/oauth/callback",  # external IdP redirects (cannot carry our token)
    "/xero/callback",
    "/auth/callback",
)


@app.before_request
def require_api_token():
    # Allow CORS preflight without auth — flask-cors handles the reply
    if request.method == "OPTIONS":
        return None
    path = request.path
    # Non-/api/ routes (frontend static assets, root) are not gated.
    if not path.startswith("/api/"):
        return None
    if path in AUTH_EXEMPT_API_PATHS:
        return None
    if any(path.startswith(p) for p in AUTH_EXEMPT_PREFIXES):
        return None
    # Accept token from Authorization header or ?token= query param (SSE).
    token = None
    auth_header = request.headers.get("Authorization", "")
    if auth_header.startswith("Bearer "):
        token = auth_header[7:]
    else:
        token = request.args.get("token")
    expected = get_setting("local_api_token", "")
    if not expected or not token or token != expected:
        return jsonify({"ok": False, "error": "Unauthorized"}), 401
    return None


# ── Helpers ────────────────────────────────────────────────────────────────────
def ok(data: Any = None, **kwargs):
    payload = {"ok": True}
    if data is not None:
        payload["data"] = data
    payload.update(kwargs)
    return jsonify(payload)


def err(message: str, status: int = 400):
    return jsonify({"ok": False, "error": message}), status


def require_loader(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if _loader is None:
            return err("Plugin loader not initialised", 503)
        return f(*args, **kwargs)
    return wrapper


# ── Health ─────────────────────────────────────────────────────────────────────
@app.route("/api/health")
def health():
    uptime_secs = int((datetime.now() - _start_time).total_seconds())
    h, rem = divmod(uptime_secs, 3600)
    m, s = divmod(rem, 60)
    return ok({
        "status": "ok",
        "uptime": f"{h}h {m}m {s}s",
        "version": get_setting("app_version", "2.4.1"),
        "fast_model": get_claude_model_fast(),
        "reasoning_model": get_claude_model_reasoning(),
    })


# ── Plugins ────────────────────────────────────────────────────────────────────
@app.route("/api/plugins")
@require_loader
def list_plugins():
    from plugin_loader import _is_plugin_allowed_in_mode, TEMPLATE_PLUGIN_IDS
    plugins = []
    states = get_all_plugin_states()
    for lp in _loader.get_plugins():
        inst = lp.instance
        # Hide template plugins — they exist as scaffolding only
        if lp.plugin_id in TEMPLATE_PLUGIN_IDS:
            continue
        # Hide plugins that don't belong to the current reception_mode
        if not _is_plugin_allowed_in_mode(inst):
            continue
        state = states.get(lp.plugin_id, {})
        plugins.append({
            "id": lp.plugin_id,
            "name": inst.name,
            "description": getattr(inst, "description", ""),
            "enabled": lp.enabled,
            "status": "disabled" if not lp.enabled else "idle",
            "lastRun": _format_dt(lp.last_run),
            "nextRun": _format_next(lp),
            "schedule": _schedule_label(lp),
            "runsToday": state.get("runs_today", 0),
            "successRate": state.get("success_rate", 100),
            "model": getattr(inst, "MODEL_TIER", "haiku"),
            "category": getattr(inst, "category", "universal"),
        })
    return ok(plugins)


@app.route("/api/plugins/<plugin_id>/enable", methods=["POST"])
@require_loader
def enable_plugin(plugin_id):
    body = request.get_json(silent=True) or {}
    enabled = body.get("enabled", True)
    _loader.set_plugin_enabled(plugin_id, enabled)
    return ok({"plugin_id": plugin_id, "enabled": enabled})


@app.route("/api/plugins/<plugin_id>/run", methods=["POST"])
@require_loader
def run_plugin(plugin_id):
    def _run():
        _loader.run_plugin(plugin_id, manual=True)
    threading.Thread(target=_run, daemon=True).start()
    return ok({"plugin_id": plugin_id, "triggered": True})


# ── Activity ───────────────────────────────────────────────────────────────────
@app.route("/api/activity")
def activity():
    limit = int(request.args.get("limit", 100))
    rows = get_recent_activity(limit)
    formatted = []
    for r in rows:
        action = r.get("action") or r.get("subject") or ""
        formatted.append({
            "id": r.get("id"),
            "time": _format_time(r.get("timestamp", "")),
            "plugin": r.get("classification") or r.get("from_email") or "",
            "action": action,
            "status": "success",
        })
    return ok(formatted)


# ── Activity SSE stream ────────────────────────────────────────────────────────
# Subscribers receive new activity rows as Server-Sent Events
_sse_subscribers: list[queue.Queue] = []
_sse_lock = threading.Lock()


def _broadcast_activity(entry: dict):
    """Push a new activity entry to all connected SSE clients."""
    with _sse_lock:
        dead = []
        for q in _sse_subscribers:
            try:
                q.put_nowait(entry)
            except queue.Full:
                dead.append(q)
        for q in dead:
            _sse_subscribers.remove(q)


# Wire into EventBus so every plugin.run.complete event triggers a broadcast
def _wire_sse_to_event_bus():
    from event_bus import EventBus
    def _on_plugin_complete(event):
        data = event.payload if hasattr(event, "payload") else {}
        entry = {
            "id": f"live-{int(time.time()*1000)}",
            "time": datetime.now().strftime("%H:%M:%S"),
            "plugin": data.get("plugin_id", "unknown"),
            "action": data.get("summary") or data.get("message", "Plugin run completed"),
            "status": "success" if data.get("success", True) else "error",
        }
        _broadcast_activity(entry)

    def _on_plugin_failed(event):
        data = event.payload if hasattr(event, "payload") else {}
        _broadcast_activity({
            "id": f"live-{int(time.time()*1000)}",
            "time": datetime.now().strftime("%H:%M:%S"),
            "plugin": data.get("plugin_id", "unknown"),
            "action": data.get("error", "Plugin run failed"),
            "status": "error",
        })

    EventBus.subscribe("plugin.run.complete", _on_plugin_complete, subscriber_id="sse_bridge")
    EventBus.subscribe("plugin.run.failed", _on_plugin_failed, subscriber_id="sse_bridge_fail")


@app.route("/api/activity/stream")
def activity_stream():
    """Server-Sent Events endpoint — streams new activity entries in real time."""
    client_q: queue.Queue = queue.Queue(maxsize=50)
    with _sse_lock:
        _sse_subscribers.append(client_q)

    def generate():
        # Send a ping immediately so the client knows the connection is live
        yield "event: ping\ndata: connected\n\n"
        try:
            while True:
                try:
                    entry = client_q.get(timeout=20)
                    yield f"data: {json.dumps(entry)}\n\n"
                except queue.Empty:
                    # Keepalive ping every 20s
                    yield "event: ping\ndata: keepalive\n\n"
        except GeneratorExit:
            pass
        finally:
            with _sse_lock:
                if client_q in _sse_subscribers:
                    _sse_subscribers.remove(client_q)

    # Only echo back the Origin header if the caller's origin is in our
    # allowlist. Same-origin (no Origin header, e.g. EventSource from our own
    # frontend) works normally — the header only matters to cross-origin
    # callers, and we refuse to grant them access to the activity stream.
    origin = request.headers.get("Origin", "")
    allow_origin = origin if origin in ALLOWED_ORIGINS else ""
    headers = {
        "Cache-Control": "no-cache",
        "X-Accel-Buffering": "no",
    }
    if allow_origin:
        headers["Access-Control-Allow-Origin"] = allow_origin
        headers["Vary"] = "Origin"
    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers=headers,
    )


# ── Approvals ──────────────────────────────────────────────────────────────────
def _approval_to_dict(a) -> dict:
    """Serialize a PendingAction dataclass to a plain JSON-safe dict."""
    status = getattr(a, "status", "")
    return {
        "action_id":     getattr(a, "action_id", None),
        "plugin_id":     getattr(a, "plugin_id", ""),
        "action_type":   getattr(a, "action_type", ""),
        "description":   getattr(a, "description", ""),
        "payload":       getattr(a, "payload", {}),
        "confidence":    getattr(a, "confidence", 0),
        "status":        status.value if hasattr(status, "value") else str(status),
        "created_at":    getattr(a, "created_at", ""),
        "expires_at":    getattr(a, "expires_at", ""),
        "reviewed_at":   getattr(a, "reviewed_at", None),
        "reviewer_note": getattr(a, "reviewer_note", None),
    }


@app.route("/api/approvals")
def list_approvals():
    if _approval_queue is None:
        return ok([])
    items = _approval_queue.list_pending()
    return ok([_approval_to_dict(a) for a in items])


@app.route("/api/approvals/<action_id>/approve", methods=["POST"])
def approve_action(action_id):
    from config import is_active_mode
    if not is_active_mode():
        return jsonify({
            "ok": False,
            "error": "Cannot approve actions in passive mode. Switch to active mode first.",
        }), 403
    if _approval_queue is None:
        return err("Approval queue not initialised", 503)
    body = request.get_json(silent=True) or {}
    reviewer_note = body.get("reviewer_note", "")
    _approval_queue.approve(int(action_id), reviewer_note=reviewer_note)
    return ok({"action_id": action_id, "approved": True})


@app.route("/api/approvals/<action_id>/reject", methods=["POST"])
def reject_action(action_id):
    from config import is_active_mode
    if not is_active_mode():
        return jsonify({
            "ok": False,
            "error": "Cannot reject actions in passive mode. Switch to active mode first.",
        }), 403
    if _approval_queue is None:
        return err("Approval queue not initialised", 503)
    body = request.get_json(silent=True) or {}
    reviewer_note = body.get("reviewer_note", "")
    _approval_queue.reject(int(action_id), reviewer_note=reviewer_note)
    return ok({"action_id": action_id, "rejected": True})


@app.route("/api/approvals/<action_id>/edit", methods=["POST"])
def edit_approval(action_id):
    """Edit the payload of a pending action before approving (edit-before-approve)."""
    if _approval_queue is None:
        return err("Approval queue not initialised", 503)
    body = request.get_json(silent=True) or {}
    updated_payload = body.get("payload")
    if updated_payload is None:
        return err("'payload' field is required")
    success = _approval_queue.edit_payload(int(action_id), updated_payload)
    if not success:
        return err("Action not found or not pending", 404)
    return ok({"action_id": action_id, "updated": True})


# ── Memory ─────────────────────────────────────────────────────────────────────
def _get_memory_store():
    """Return the live MemoryStore owned by the PluginLoader, or None.

    The loader creates a single MemoryStore in its __init__; sharing that
    instance here avoids a second ChromaDB PersistentClient pointed at the
    same on-disk directory (which fails on Windows with a file-lock error).
    """
    if _loader is not None and getattr(_loader, "_memory", None) is not None:
        return _loader._memory
    return None


@app.route("/api/memory")
def list_memory():
    """List memory records.

    - If `q` is given → semantic search across client interactions.
    - If `q` is empty or missing → return the most recent interactions.
    """
    try:
        query = (request.args.get("q") or "").strip()
        limit = int(request.args.get("limit", 50))
        store = _get_memory_store()
        if store is None:
            return jsonify({"ok": True, "data": [], "total": 0, "degraded": True}), 200

        if query:
            results = store.search(query, n_results=limit, collection="interactions")
        else:
            try:
                raw = store._interactions.get(limit=limit)
                results = []
                if raw and raw.get("ids"):
                    ids = raw["ids"]
                    docs = raw.get("documents") or []
                    metas = raw.get("metadatas") or []
                    for i, doc_id in enumerate(ids):
                        results.append({
                            "id": doc_id,
                            "content": docs[i] if i < len(docs) else "",
                            "metadata": metas[i] if i < len(metas) else {},
                            "distance": 0.0,
                            "relevance": None,
                        })
                    results.sort(
                        key=lambda x: x.get("metadata", {}).get("epoch", 0),
                        reverse=True,
                    )
            except Exception as e:
                logger.error(f"Memory fetch failed: {e}", exc_info=True)
                results = []

        return jsonify({
            "ok": True,
            "data": results,
            "total": store.count("interactions"),
        })
    except Exception as e:
        logger.error(f"API error in {request.path}: {e}", exc_info=True)
        return jsonify({
            "ok": True, "data": [], "degraded": True, "error": str(e),
        }), 200


@app.route("/api/memory/search", methods=["POST"])
def search_memory():
    """Advanced semantic search with optional metadata filters."""
    try:
        body = request.json or {}
        query = body.get("query", "")
        collection = body.get("collection", "interactions")
        filters = body.get("filters")
        limit = int(body.get("limit", 10))

        store = _get_memory_store()
        if store is None:
            return jsonify({"ok": True, "data": []})

        results = store.search(
            query, n_results=limit, filters=filters, collection=collection,
        )
        return jsonify({"ok": True, "data": results})
    except Exception as e:
        logger.error(f"API error in {request.path}: {e}", exc_info=True)
        return err(str(e))


@app.route("/api/memory/stats")
def memory_stats():
    """Counts for each semantic memory collection."""
    store = _get_memory_store()
    if store is None:
        return jsonify({
            "ok": True, "interactions": 0, "lessons": 0, "documents": 0,
        })
    try:
        return jsonify({
            "ok": True,
            "interactions": store.count("interactions"),
            "lessons": store.count("lessons"),
            "documents": store.count("documents"),
        })
    except Exception as e:
        logger.error(f"API error in {request.path}: {e}", exc_info=True)
        return jsonify({
            "ok": True, "interactions": 0, "lessons": 0, "documents": 0,
            "degraded": True, "error": str(e),
        })


@app.route("/api/memory/<record_id>", methods=["DELETE"])
def delete_memory(record_id):
    try:
        store = _get_memory_store()
        if store is None:
            return err("Memory store unavailable")
        store.delete(record_id)
        return ok({"deleted": record_id})
    except Exception as e:
        return err(str(e))


# ── Events ─────────────────────────────────────────────────────────────────────
@app.route("/api/events")
def list_events():
    limit = int(request.args.get("limit", 50))
    history = EventBus.get_history(limit=limit)
    formatted = []
    for evt in reversed(history):
        ts = getattr(evt, "timestamp", 0)
        try:
            time_str = datetime.fromtimestamp(float(ts)).strftime("%H:%M:%S")
        except Exception:
            time_str = ""
        formatted.append({
            "id": f"evt-{int(float(ts)*1000)}" if ts else "",
            "time": time_str,
            "type": getattr(evt, "type", ""),
            "source": getattr(evt, "source", ""),
            "payload": str(getattr(evt, "payload", "")),
        })
    return ok(formatted)


# ── KPI ────────────────────────────────────────────────────────────────────────
@app.route("/api/kpi")
def kpi():
    try:
        from kpi_monitor import get_kpi_config, get_recent_alerts
        configs = get_kpi_config()
        alerts = get_recent_alerts(limit=50)
        latest_by_kpi = {}
        for a in alerts:
            kid = a.get("kpi_id")
            if kid and kid not in latest_by_kpi:
                latest_by_kpi[kid] = a
        metrics = []
        for cfg in configs:
            kid = cfg.get("kpi_id")
            latest = latest_by_kpi.get(kid, {})
            metrics.append({
                "kpi_id":     kid,
                "label":      cfg.get("label", kid),
                "description":cfg.get("description", ""),
                "threshold":  cfg.get("threshold", 0),
                "enabled":    bool(cfg.get("enabled", 1)),
                "severity":   cfg.get("severity", "warning"),
                "unit":       cfg.get("unit", ""),
                "value":      latest.get("value"),
                "message":    latest.get("message", ""),
                "last_alert": latest.get("timestamp", ""),
            })
        return ok(metrics)
    except Exception as e:
        logger.error(f"API error in {request.path}: {e}", exc_info=True)
        return jsonify({
            "ok": True, "data": [], "degraded": True, "error": str(e),
        }), 200


# ── Usage ──────────────────────────────────────────────────────────────────────
@app.route("/api/usage")
def usage():
    try:
        s = get_usage_summary(days=30)
        return ok({
            "todayCost": s.get("today_cost_aud", 0),
            "monthlyCost": s.get("this_month_cost_aud", 0),
            "monthlyBudget": float(get_setting("monthly_ai_budget_aud", "100")),
            "totalCalls": s.get("total_calls", 0),
            "totalTokensIn": s.get("total_tokens", 0),
            "totalTokensOut": 0,
            "byPlugin": [
                {"plugin_id": r["plugin_id"], "calls": r["calls"],
                 "tokens": r["tokens"], "cost_aud": r["cost_aud"]}
                for r in (s.get("by_plugin") or [])
            ],
            "byDay": [
                {"date": r["date"], "calls": r["calls"],
                 "tokens": r["tokens"], "cost_aud": r["cost_aud"]}
                for r in (s.get("by_day") or [])
            ],
        })
    except Exception as e:
        logger.error(f"API error in {request.path}: {e}", exc_info=True)
        return jsonify({
            "ok": True,
            "data": {"todayCost": 0, "monthlyCost": 0, "monthlyBudget": 100,
                     "totalCalls": 0, "totalTokensIn": 0, "totalTokensOut": 0,
                     "byPlugin": [], "byDay": []},
            "degraded": True, "error": str(e),
        }), 200


# ── Settings ───────────────────────────────────────────────────────────────────
@app.route("/api/settings")
def get_settings():
    s = get_all_settings()
    # Mask sensitive keys
    for key in ("anthropic_api_key", "fusesign_api_key", "teams_webhook_url",
                "statementhub_api_key",
                "xero_client_secret", "xero_access_token", "xero_refresh_token"):
        if s.get(key):
            s[key] = s[key][:4] + "••••••••••••••••••••"
    s["fast_model"] = get_claude_model_fast()
    s["reasoning_model"] = get_claude_model_reasoning()
    s["opus_model"] = get_setting("opus_model", "claude-opus-4-6")
    # Add Xero OAuth status
    try:
        from xero_oauth import is_configured as xero_is_configured, is_authorised as xero_is_authorised
        s["xero_configured"] = xero_is_configured()
        s["xero_authorised"] = xero_is_authorised()
    except Exception:
        s["xero_configured"] = False
        s["xero_authorised"] = False
    return ok(s)


@app.route("/api/settings", methods=["POST"])
def save_settings():
    body = request.get_json(silent=True) or {}
    # Only save non-masked values
    safe_keys = {
        "anthropic_api_key", "outlook_email",
        "fusesign_api_key", "teams_webhook_url",
        "statementhub_api_key", "statementhub_base_url",
        "confidence_threshold", "heartbeat_interval_seconds",
        "draft_mode", "auto_update_enabled",
        "fast_model", "reasoning_model", "opus_model",
        "monthly_ai_budget_aud",
        "skip_public_holidays", "public_holiday_state",
        "reception_mode", "staff_profile",
        "email_signature",
        # Xero OAuth credentials
        "xero_client_id", "xero_client_secret",
    }
    saved = []
    for key, value in body.items():
        if key in safe_keys and "••••" not in str(value):
            set_setting(key, str(value))
            saved.append(key)
    # Re-detect models if API key changed
    if "anthropic_api_key" in saved:
        try:
            api_key = get_setting("anthropic_api_key", "")
            if api_key:
                update_claude_models(api_key)
        except Exception:
            pass
    return ok({"saved": saved})


# ── Signature Image ──────────────────────────────────────────────────────────
# Uploaded signature image lives at DATA_DIR/signature.png. All uploads are
# normalised to PNG via Pillow so graph_client has a single well-known path
# and format to read when embedding the image inline in outgoing drafts.
SIGNATURE_IMAGE_PATH = config.DATA_DIR / "signature.png"
ALLOWED_SIGNATURE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif"}


@app.route("/api/settings/signature-image", methods=["POST"])
def upload_signature_image():
    """Accept a single image upload, convert to PNG, save to DATA_DIR/signature.png."""
    if "file" not in request.files:
        return err("No file provided", 400)

    file = request.files["file"]
    if not file.filename:
        return err("No filename", 400)

    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_SIGNATURE_EXTENSIONS:
        return err(f"File type {ext} not supported (use .png, .jpg, .gif)", 400)

    try:
        from PIL import Image
        import io
        data = file.read()
        if len(data) > MAX_UPLOAD_SIZE:
            return err("File too large (max 25MB)", 413)
        img = Image.open(io.BytesIO(data))
        # Flatten transparency onto white? No — preserve transparency by saving
        # as PNG RGBA. Convert palette/other modes to RGBA so PNG save succeeds.
        if img.mode not in ("RGB", "RGBA"):
            img = img.convert("RGBA")
        SIGNATURE_IMAGE_PATH.parent.mkdir(parents=True, exist_ok=True)
        img.save(str(SIGNATURE_IMAGE_PATH), format="PNG")
    except Exception as e:
        return err(f"Failed to process image: {e}", 400)

    return ok({"path": str(SIGNATURE_IMAGE_PATH), "size": SIGNATURE_IMAGE_PATH.stat().st_size})


@app.route("/api/settings/signature-image", methods=["GET"])
def get_signature_image():
    """Return the saved signature image file, or 404 if none."""
    if not SIGNATURE_IMAGE_PATH.exists():
        return err("No signature image", 404)
    return send_file(str(SIGNATURE_IMAGE_PATH), mimetype="image/png")


@app.route("/api/settings/signature-image", methods=["DELETE"])
def delete_signature_image():
    """Remove the saved signature image."""
    try:
        SIGNATURE_IMAGE_PATH.unlink(missing_ok=True)
    except Exception as e:
        return err(f"Failed to remove: {e}", 500)
    return ok({"removed": True})


@app.route("/api/settings/test/<service>", methods=["POST"])
def test_connection(service):
    try:
        from gateway_client import GatewayClient
        gw = GatewayClient()
        gw.load()
        if service == "xpm":
            result = gw.xpm.list_clients(page=1, page_size=1)
            return ok({"connected": True, "service": "xpm"})
        elif service == "fusesign":
            result = gw.fusesign.list_envelopes(page_size=1)
            return ok({"connected": True, "service": "fusesign"})
        elif service == "teams":
            gw.teams.send_alert(title="CoWorker", body="Connection test successful ✅")
            return ok({"connected": True, "service": "teams"})
        else:
            return err(f"Unknown service: {service}")
    except Exception as e:
        return ok({"connected": False, "service": service, "error": str(e)})


# ── Xero OAuth ─────────────────────────────────────────────────────────────────

# Background thread state for OAuth flow
_xero_auth_thread: threading.Thread | None = None
_xero_auth_status: dict = {"status": "idle", "message": ""}


@app.route("/api/xero/status")
def xero_status():
    """Return current Xero OAuth status."""
    try:
        from xero_oauth import is_configured, is_authorised, get_tenant_id
        return ok({
            "configured":  is_configured(),
            "authorised":  is_authorised(),
            "tenant_id":   get_tenant_id(),
            "client_id":   get_setting("xero_client_id", ""),
            "auth_status": _xero_auth_status,
        })
    except Exception as e:
        return err(str(e))


@app.route("/api/xero/start-auth", methods=["POST"])
def xero_start_auth():
    """
    Start the Xero OAuth Authorization Code flow in a background thread.
    Opens the user's browser to the Xero login page.
    The callback is handled by /oauth/callback below.
    """
    global _xero_auth_thread, _xero_auth_status

    if _xero_auth_thread and _xero_auth_thread.is_alive():
        return ok({"status": "in_progress", "message": "OAuth flow already in progress"})

    _xero_auth_status = {"status": "in_progress", "message": "Opening Xero login..."}

    def _run_auth():
        global _xero_auth_status
        try:
            from xero_oauth import start_auth_flow
            token_data = start_auth_flow(timeout=300)
            _xero_auth_status = {
                "status":  "success",
                "message": "Xero connected successfully!",
                "scope":   token_data.get("scope", ""),
            }
        except Exception as e:
            _xero_auth_status = {"status": "error", "message": str(e)}

    _xero_auth_thread = threading.Thread(target=_run_auth, daemon=True)
    _xero_auth_thread.start()

    return ok({"status": "in_progress", "message": "Xero login page opened in browser"})


@app.route("/oauth/callback")
def xero_oauth_callback():
    """
    Xero OAuth callback — receives the authorization code after user login.
    Passes the code to the waiting start_auth_flow() thread.
    """
    code  = request.args.get("code", "")
    state = request.args.get("state", "")
    error = request.args.get("error", "")

    try:
        from xero_oauth import notify_oauth_callback
        notify_oauth_callback(code=code, state=state, error=error)
    except Exception as e:
        return f"<h1>Error</h1><p>{html.escape(str(e))}</p>", 500

    if error:
        safe_error = html.escape(error)
        return (
            "<!DOCTYPE html><html><head><title>Xero — Error</title>"
            "<style>body{font-family:sans-serif;text-align:center;padding:60px}</style></head>"
            f"<body><h1>\u274c Xero Connection Failed</h1><p>{safe_error}</p>"
            "<p>You can close this window.</p></body></html>"
        ), 400

    return (
        "<!DOCTYPE html><html><head><title>MCS CoWorker — Xero Connected</title>"
        "<style>body{font-family:sans-serif;text-align:center;padding:60px;background:#f0f4f8}"
        "h1{color:#13b5ea}p{color:#444}</style></head>"
        "<body><h1>&#10003; Xero Connected!</h1>"
        "<p>MCS CoWorker is now connected to Xero XPM.</p>"
        "<p>You can close this window and return to CoWorker.</p></body></html>"
    )


@app.route("/api/xero/disconnect", methods=["POST"])
def xero_disconnect():
    """Revoke the Xero refresh token and clear all stored tokens."""
    try:
        from xero_oauth import revoke_token
        revoke_token()
        return ok({"disconnected": True})
    except Exception as e:
        return err(str(e))


# ── Chat ───────────────────────────────────────────────────────────────────────
@app.route("/api/chat/upload", methods=["POST"])
def upload_chat_file():
    """Upload a file for use in the current chat conversation.

    Stored under DATA_DIR / "chat_uploads" keyed by a random id. The returned
    `{id, name, type, size}` payload is what the frontend should echo back in
    the `files` array of the next POST to /api/chat.
    """
    if "file" not in request.files:
        return jsonify({"ok": False, "error": "No file provided"}), 400

    file = request.files["file"]
    if not file.filename:
        return jsonify({"ok": False, "error": "No filename"}), 400

    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_UPLOAD_EXTENSIONS:
        return jsonify({"ok": False, "error": f"File type {ext} not supported"}), 400

    file_id = uuid.uuid4().hex[:12]
    safe_name = f"{file_id}{ext}"
    file_path = CHAT_UPLOAD_DIR / safe_name
    file.save(str(file_path))

    size = file_path.stat().st_size
    if size > MAX_UPLOAD_SIZE:
        file_path.unlink(missing_ok=True)
        return jsonify({"ok": False, "error": "File too large (max 25MB)"}), 413

    return jsonify({
        "ok": True,
        "file": {
            "id": file_id,
            "name": file.filename,
            "path": str(file_path),
            "size": size,
            "type": ext,
        },
    })


def _build_file_content_blocks(file_refs: list) -> list:
    """Convert frontend file references into Anthropic content blocks.

    PDFs → `document` blocks (Claude's native PDF understanding).
    Images → `image` blocks (base64).
    Spreadsheets / CSV / Word / text → inline `text` blocks wrapped in
    <uploaded_file name="..."> tags.

    Capped at MAX_UPLOADS_PER_MESSAGE and 100 KB of extracted text per file to
    bound token spend per call.
    """
    blocks: list = []
    for file_ref in (file_refs or [])[:MAX_UPLOADS_PER_MESSAGE]:
        file_id = file_ref.get("id")
        ext = (file_ref.get("type") or "").lower()
        if not file_id or not ext:
            continue

        file_path = CHAT_UPLOAD_DIR / f"{file_id}{ext}"
        if not file_path.exists():
            logger.warning(f"Chat file ref missing on disk: {file_path}")
            continue

        name = file_ref.get("name", file_path.name)

        try:
            if ext == ".pdf":
                with open(file_path, "rb") as f:
                    b64 = base64.b64encode(f.read()).decode("ascii")
                blocks.append({
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": b64,
                    },
                })

            elif ext in (".jpg", ".jpeg", ".png", ".gif"):
                with open(file_path, "rb") as f:
                    b64 = base64.b64encode(f.read()).decode("ascii")
                media_type = mimetypes.guess_type(f"file{ext}")[0] or "image/png"
                blocks.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": b64,
                    },
                })

            elif ext == ".csv":
                content = file_path.read_text(encoding="utf-8", errors="replace")
                blocks.append({
                    "type": "text",
                    "text": f'<uploaded_file name="{name}">\n{content[:100000]}\n</uploaded_file>',
                })

            elif ext == ".xlsx":
                import openpyxl
                wb = openpyxl.load_workbook(str(file_path), data_only=True)
                sheets = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    rows = []
                    for row in ws.iter_rows(values_only=True):
                        rows.append("\t".join(str(c) if c is not None else "" for c in row))
                    sheets.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
                content = "\n\n".join(sheets)
                blocks.append({
                    "type": "text",
                    "text": f'<uploaded_file name="{name}">\n{content[:100000]}\n</uploaded_file>',
                })

            elif ext == ".docx":
                from docx import Document as DocxDocument
                doc = DocxDocument(str(file_path))
                content = "\n".join(p.text for p in doc.paragraphs)
                blocks.append({
                    "type": "text",
                    "text": f'<uploaded_file name="{name}">\n{content[:100000]}\n</uploaded_file>',
                })

            elif ext == ".txt":
                content = file_path.read_text(encoding="utf-8", errors="replace")
                blocks.append({
                    "type": "text",
                    "text": f'<uploaded_file name="{name}">\n{content[:100000]}\n</uploaded_file>',
                })

        except Exception as e:
            logger.error(f"Failed to read {ext} file {name}: {e}", exc_info=True)

    return blocks


@app.route("/api/agents", methods=["GET"])
def list_agents():
    """Return all available specialist agents for the Chat tab dropdown."""
    agents = get_all_agents()
    return jsonify({
        "ok": True,
        "agents": [
            {
                "id": a.id,
                "name": a.name,
                "description": a.description,
                "icon": a.icon,
                "category": a.category,
                "supports_files": a.supports_files,
                "file_types": a.file_types,
                "model_preference": a.model_preference,
            }
            for a in agents.values()
        ],
    })


@app.route("/api/chat", methods=["POST"])
def chat():
    if not _check_chat_rate():
        return jsonify({
            "ok": False,
            "error": f"Rate limit exceeded. Max {CHAT_RATE_LIMIT} requests per {CHAT_RATE_WINDOW} seconds."
        }), 429

    body = request.get_json(silent=True) or {}
    messages = body.get("messages", [])
    if not messages:
        return err("No messages provided")

    # Resolve which specialist agent is answering this turn. Default is the
    # plugin builder so existing Chat behaviour is preserved.
    agent_id = body.get("agent_id", "plugin_builder")
    agent = get_agent(agent_id)
    if not agent:
        return jsonify({"ok": False, "error": f"Unknown agent: {agent_id}"}), 400

    # Optional client context — supplied by the Chat UI when the user has
    # nominated which client this conversation is about.
    client_name = body.get("client_name") or None
    entity_name = body.get("entity_name") or None

    # Cap per-message size to bound token cost per call.
    for m in messages:
        content = m.get("content", "")
        if isinstance(content, str) and len(content) > CHAT_MAX_MESSAGE_LEN:
            return jsonify({
                "ok": False,
                "error": f"Message too long. Max {CHAT_MAX_MESSAGE_LEN} characters."
            }), 413

    # Truncate history to the most recent N turns so runaway conversations
    # don't balloon the context window every call.
    if len(messages) > CHAT_MAX_HISTORY:
        messages = messages[-CHAT_MAX_HISTORY:]

    # Build file content blocks from any attachments referenced in this call.
    attached_files = body.get("files", []) or []
    file_content_blocks = _build_file_content_blocks(attached_files)

    # If there are file blocks, attach them to the LAST user message by
    # converting its string content into a list of blocks: [files..., text].
    if file_content_blocks and messages:
        last = messages[-1]
        if last.get("role") == "user":
            orig = last.get("content", "")
            text = orig if isinstance(orig, str) else ""
            last["content"] = file_content_blocks + [{"type": "text", "text": text}]

    try:
        import anthropic as anthropic_lib
        api_key = get_setting("anthropic_api_key", "")
        if not api_key:
            return err("Anthropic API key not configured")

        client = anthropic_lib.Anthropic(api_key=api_key)

        # Pull the raw text of the last user message. Used below for both
        # tier detection (plugin_builder only) and memory retrieval.
        raw_last = messages[-1].get("content", "")
        if isinstance(raw_last, list):
            last_text = " ".join(
                b.get("text", "") for b in raw_last if isinstance(b, dict) and b.get("type") == "text"
            )
        else:
            last_text = raw_last

        # Model routing:
        #   - plugin_builder keeps the legacy tier-keyword scan so heavy
        #     plugin tasks (XPM, Teams, KPIs, etc.) escalate to the reasoning
        #     model while simple template jobs stay on Haiku.
        #   - Every other specialist uses its configured model preference.
        if agent.id == "plugin_builder":
            tier2_keywords = ["xpm", "fusesign", "teams", "memory", "workflow",
                              "report", "wip", "debtor", "engagement", "onboard",
                              "gateway", "event", "heartbeat", "kpi"]
            is_tier2 = any(k in last_text.lower() for k in tier2_keywords)
            model = get_claude_model_reasoning() if is_tier2 else get_claude_model_fast()
            max_tokens = 4096 if is_tier2 else 2048
        else:
            if agent.model_preference == "opus":
                model = get_setting("opus_model") or "claude-opus-4-6"
                max_tokens = 4096
            elif agent.model_preference == "sonnet":
                model = get_claude_model_reasoning()
                max_tokens = 4096
            else:
                model = get_claude_model_fast()
                max_tokens = 2048

        # System prompt:
        #   - plugin_builder retains its live-data block (pending approvals,
        #     ASIC, debtors, recent activity) so the builder can answer
        #     "what's going on in the practice" questions.
        #   - specialists use their file-loaded prompt verbatim.
        if agent.id == "plugin_builder":
            system_prompt = _build_chat_system_prompt()
        else:
            system_prompt = agent.system_prompt

        # Tell a file-aware specialist to actually look at the uploads.
        if agent.supports_files and file_content_blocks:
            system_prompt += (
                "\n\nThe user has uploaded files with this message. "
                "Analyse them according to your specialist instructions."
            )

        # Inject relevant semantic memory for non-plugin_builder specialists.
        # Memory contains client interactions, prior advice, and lessons —
        # useful for tax specialists that need practice-specific context.
        if agent.id != "plugin_builder":
            store = _get_memory_store()
            if store is not None:
                try:
                    if client_name:
                        from client_utils import normalise_client_name
                        normalised_for_lookup = normalise_client_name(client_name)
                        client_context = store.format_for_prompt(
                            f"{normalised_for_lookup} {last_text[:100]}",
                            n_results=5,
                        )
                        if client_context:
                            system_prompt += (
                                f"\n\nRelevant client history from practice memory:\n{client_context}"
                            )
                    elif last_text:
                        memory_context = store.format_for_prompt(last_text, n_results=5)
                        if memory_context:
                            system_prompt += f"\n\n{memory_context}"
                except Exception as e:
                    logger.warning(f"Memory context injection skipped: {e}")

        # Specialists whose prompts instruct them to pull live data
        # (Fair Work, ATO, ASIC, etc.) need the Claude web search tool to
        # avoid fabricating rates from training data.
        AGENTS_NEEDING_SEARCH = {
            "gst", "smsf", "div7a", "trusts", "tax_structure",
            "payroll", "individual_tax", "general",
        }
        create_kwargs = {
            "model": model,
            "max_tokens": max_tokens,
            "system": system_prompt,
            "messages": messages,
        }
        if agent.id in AGENTS_NEEDING_SEARCH:
            create_kwargs["tools"] = [{
                "type": "web_search_20250305",
                "name": "web_search",
            }]

        response = client.messages.create(**create_kwargs)
        # With web search, content may include tool_use / search_result blocks
        # alongside text. Concatenate every text block in order.
        response_text = ""
        for block in response.content:
            text = getattr(block, "text", None)
            if text:
                response_text += text
        if not response_text and response.content:
            # Fallback for unexpected content shapes.
            first = response.content[0]
            response_text = getattr(first, "text", "") or ""

        # Persist the exchange under the client so future turns and other
        # plugins (Smart Responder etc.) can recall the advice given.
        if client_name and agent.id != "plugin_builder":
            try:
                from client_utils import normalise_client_name
                store = _get_memory_store()
                if store is not None:
                    normalised = normalise_client_name(client_name)
                    summary = (
                        f"[{agent.name}] Q: {last_text[:200]} → A: {response_text[:300]}"
                    )
                    store.store_client_interaction(
                        client_name=normalised,
                        entity_name=entity_name,
                        interaction_type="ai_chat",
                        summary=summary,
                        metadata={
                            "agent_id": agent.id,
                            "agent_name": agent.name,
                            "full_question": last_text[:1000],
                        },
                    )
            except Exception as e:
                logger.warning(f"AI chat memory write skipped: {e}")

        return ok({
            "content": response_text,
            "model": model,
            "agent_id": agent.id,
            "agent_name": agent.name,
            "usage": {
                "input_tokens": response.usage.input_tokens,
                "output_tokens": response.usage.output_tokens,
            },
            "files_processed": len(file_content_blocks),
        })
    except Exception as e:
        return err(f"Chat error: {str(e)}")


@app.route("/api/clients/names", methods=["GET"])
def get_client_names():
    """Return all known client names for autocomplete in the Chat UI."""
    store = _get_memory_store()
    if store is None:
        return jsonify({"ok": True, "clients": []})
    try:
        names = store.get_all_clients()
        return jsonify({"ok": True, "clients": names})
    except Exception as e:
        logger.error(f"get_client_names failed: {e}", exc_info=True)
        return jsonify({"ok": True, "clients": []})


# ── System ─────────────────────────────────────────────────────────────────────
@app.route("/api/system/status")
def system_status():
    uptime_secs = int((datetime.now() - _start_time).total_seconds())
    h, rem = divmod(uptime_secs, 3600)
    m, _ = divmod(rem, 60)
    try:
        store = _get_memory_store()
        mem_count = store.count() if store is not None else 0
    except Exception:
        mem_count = 0
    try:
        from token_meter import get_usage_summary
        cost_today = f"${get_usage_summary().get('today_cost_aud', 0):.2f}"
    except Exception:
        cost_today = "$0.00"
    try:
        from event_bus import EventBus
        tick = len(EventBus.get_history(limit=10000))
    except Exception:
        tick = 0

    return ok({
        "heartbeat": "Active" if _loader and _loader._running else "Stopped",
        "heartbeatTick": tick,
        "fastModel": get_claude_model_fast(),
        "reasoningModel": get_claude_model_reasoning(),
        "memoryRecords": mem_count,
        "costToday": cost_today,
        "uptime": f"{h}h {m}m",
        "version": get_setting("app_version", "2.4.1"),
        "updateAvailable": False,
        "offlineMode": get_setting("_offline_mode", "false") == "true",
        "processing_mode": get_setting("processing_mode", "active"),
    })


@app.route("/api/processing-mode", methods=["GET"])
def get_processing_mode():
    """Return the current per-machine processing mode."""
    from config import is_active_mode
    return jsonify({
        "ok": True,
        "mode": "active" if is_active_mode() else "passive",
    })


@app.route("/api/processing-mode", methods=["POST"])
def set_processing_mode():
    """Set the per-machine processing mode. Active runs plugins; passive
    is monitor-only so a second machine on the same mailbox doesn't draft
    duplicates."""
    body = request.get_json(silent=True) or {}
    mode = (body.get("mode") or "active").lower()
    if mode not in ("active", "passive"):
        return jsonify({
            "ok": False,
            "error": "Invalid mode. Use 'active' or 'passive'.",
        }), 400
    from config import save_setting
    save_setting("processing_mode", mode)
    return jsonify({"ok": True, "mode": mode})


# ── Helpers ────────────────────────────────────────────────────────────────────
def _format_dt(dt) -> str:
    if dt is None:
        return "Never"
    if isinstance(dt, datetime):
        now = datetime.now()
        diff = now - dt
        if diff.total_seconds() < 60:
            return "Just now"
        if diff.total_seconds() < 3600:
            return f"{int(diff.total_seconds() // 60)} min ago"
        if diff.total_seconds() < 86400:
            return f"{int(diff.total_seconds() // 3600)} hr ago"
        return dt.strftime("%d %b")
    return str(dt)


def _format_time(ts: str) -> str:
    try:
        dt = datetime.fromisoformat(ts)
        return dt.strftime("%H:%M")
    except Exception:
        return ts[:5] if ts else ""


def _format_next(lp) -> str:
    if not lp.enabled:
        return "Disabled"
    if lp.schedule_seconds == 0 and not lp.instance.default_schedule.is_calendar_based():
        return "On event"
    if lp._next_run_at <= 0:
        return "Soon"
    diff = lp._next_run_at - time.time()
    if diff <= 0:
        return "Now"
    if diff < 60:
        return f"{int(diff)}s"
    if diff < 3600:
        return f"{int(diff // 60)} min"
    return f"{int(diff // 3600)} hr"


def _schedule_label(lp) -> str:
    s = lp.schedule_seconds
    if s == 0:
        return "Event-driven"
    if s < 3600:
        return f"Every {s // 60} minutes"
    if s < 86400:
        return f"Every {s // 3600} hours"
    return "Daily"


def _build_chat_system_prompt() -> str:
    style = config.get_style_preferences() if hasattr(config, "get_style_preferences") else ""

    # Inject live practice context so the assistant can answer operational questions
    context_lines = []
    try:
        pending_count = _approval_queue.count_pending() if _approval_queue else 0
        context_lines.append(f"Pending approvals in queue: {pending_count}")
        if _approval_queue and pending_count > 0:
            pending_items = _approval_queue.list_pending()[:5]
            for item in pending_items:
                desc = getattr(item, "description", "") or getattr(item, "action_type", "?")
                plugin = getattr(item, "plugin_id", "?")
                context_lines.append(f"  - [{plugin}] {desc[:80]}")
    except Exception:
        pass
    try:
        from plugins.plugin_asic_returns import get_asic_returns
        asic_open = [r for r in get_asic_returns(limit=200)
                     if r.get("status") not in ("completed", "cancelled")]
        context_lines.append(f"Open ASIC annual returns: {len(asic_open)}")
        awaiting_solvency = [r for r in asic_open if r.get("status") == "awaiting_solvency"]
        overdue_asic = [r for r in asic_open if r.get("status") == "overdue"]
        if awaiting_solvency:
            context_lines.append(f"  Awaiting solvency resolution: {len(awaiting_solvency)} "
                                  f"({', '.join(r.get('company_name','?') for r in awaiting_solvency[:3])})")
        if overdue_asic:
            context_lines.append(f"  Overdue ASIC returns: {len(overdue_asic)} "
                                  f"({', '.join(r.get('company_name','?') for r in overdue_asic[:3])})")
        elif asic_open:
            names = ", ".join(r.get("company_name", "?") for r in asic_open[:5])
            context_lines.append(f"  Companies: {names}{'...' if len(asic_open) > 5 else ''}")
    except Exception:
        pass
    try:
        lessons = get_active_lessons()
        if lessons:
            lesson_text = "; ".join(l["lesson"] for l in lessons[:10])
            context_lines.append(f"Learned preferences: {lesson_text}")
    except Exception:
        pass
    try:
        recent = get_recent_activity(limit=10)
        if recent:
            activity_text = "; ".join(
                f"{r.get('classification','?')}: {r.get('subject','?')[:40]}"
                for r in recent
            )
            context_lines.append(f"Recent activity (last 10): {activity_text}")
    except Exception:
        pass
    try:
        # Overdue debtors summary from XPM invoices
        from gateway_client import XPMClient
        xpm = XPMClient()
        if xpm.is_configured and xpm.is_authorised:
            summary = xpm.get_debtor_summary()
            total = summary.get("total_outstanding", 0)
            count = summary.get("invoice_count", 0)
            over90 = summary.get("90_plus", 0)
            context_lines.append(
                f"Debtor summary: {count} outstanding invoices totalling ${total:,.0f} "
                f"(${over90:,.0f} is 90+ days overdue)"
            )
    except Exception:
        pass

    practice_context = ("\n".join(context_lines)) if context_lines else ""
    practice_name = get_setting("practice_name", "MC & S")

    return f"""You are the AI assistant built into {practice_name} CoWorker, an intelligent automation platform for an Australian accounting firm.
You can answer questions about the practice's current state, explain what plugins do, help build new automation, and advise on workflow improvements.

CURRENT PRACTICE STATE
{practice_context if practice_context else '(No live data available yet)'}

PLUGIN DEVELOPMENT
TIER 1 — Template Builder: for common email/auto-reply patterns
TIER 2 — Custom Plugin Writer: full Python plugins using:
  - context.claude_fast / context.claude_reason (dual Claude models)
  - context.memory (ChromaDB vector store)
  - context.event_bus (publish/subscribe events)
  - context.gateway.xpm (XPM practice management)
  - context.gateway.fusesign (document signing)
  - context.gateway.teams (Teams notifications)
  - context.approval_queue (confidence-based human review)

Always produce working Python code. Use PluginResult(success=True/False, message="...").
{f'Style preferences: {style}' if style else ''}"""


# ── ASIC Tracker ─────────────────────────────────────────────────────────────
@app.route("/api/asic")
def list_asic_returns():
    """List all ASIC annual return records."""
    try:
        from plugins.plugin_asic_returns import get_asic_returns
        status = request.args.get("status", None)
        limit  = int(request.args.get("limit", 100))
        rows = get_asic_returns(status=status, limit=limit)
        return ok(rows)
    except Exception as e:
        return err(str(e))


@app.route("/api/asic/<int:return_id>/mark-paid", methods=["POST"])
def asic_mark_paid(return_id):
    """Mark an ASIC return as paid and update status to 'completed' if solvency also signed."""
    try:
        from plugins.plugin_asic_returns import update_asic_return, get_asic_returns
        rows = get_asic_returns(limit=1000)
        record = next((r for r in rows if r["id"] == return_id), None)
        if not record:
            return err("ASIC return not found", 404)
        new_status = "completed" if record.get("solvency_signed") else "awaiting_solvency"
        update_asic_return(return_id, asic_paid=1, status=new_status)
        return ok({"id": return_id, "asic_paid": True, "status": new_status})
    except Exception as e:
        return err(str(e))


@app.route("/api/asic/<int:return_id>/mark-solvency-signed", methods=["POST"])
def asic_mark_solvency_signed(return_id):
    """Mark an ASIC return's solvency resolution as signed and update status."""
    try:
        from plugins.plugin_asic_returns import update_asic_return, get_asic_returns
        rows = get_asic_returns(limit=1000)
        record = next((r for r in rows if r["id"] == return_id), None)
        if not record:
            return err("ASIC return not found", 404)
        new_status = "completed" if record.get("asic_paid") else "awaiting_payment"
        update_asic_return(return_id, solvency_signed=1, status=new_status)
        return ok({"id": return_id, "solvency_signed": True, "status": new_status})
    except Exception as e:
        return err(str(e))


@app.route("/api/asic/<int:return_id>/notes", methods=["POST"])
def asic_update_notes(return_id):
    """Update the notes field on an ASIC return."""
    try:
        from plugins.plugin_asic_returns import update_asic_return
        body = request.get_json(silent=True) or {}
        notes = body.get("notes", "")
        update_asic_return(return_id, notes=notes)
        return ok({"id": return_id, "notes": notes})
    except Exception as e:
        return err(str(e))


# ── Email Rules ───────────────────────────────────────────────────────────────
@app.route("/api/rules")
def list_rules():
    return ok(get_rules())

@app.route("/api/rules", methods=["POST"])
def create_rule():
    data = request.get_json(force=True)
    save_rule(data)
    return ok(get_rules())

@app.route("/api/rules/<int:rule_id>", methods=["DELETE"])
def remove_rule(rule_id):
    delete_rule(rule_id)
    return ok()

# ── Staff ─────────────────────────────────────────────────────────────────────
@app.route("/api/staff")
def list_staff():
    return ok(get_staff())

@app.route("/api/staff", methods=["POST"])
def create_staff():
    data = request.get_json(force=True)
    save_staff(data)
    return ok(get_staff())

@app.route("/api/staff/<int:staff_id>", methods=["DELETE"])
def remove_staff(staff_id):
    delete_staff(staff_id)
    return ok()

# ── Links & Forms ─────────────────────────────────────────────────────────────
@app.route("/api/links")
def list_links():
    return ok(get_links())

@app.route("/api/links", methods=["POST"])
def create_link():
    data = request.get_json(force=True)
    save_link(data)
    return ok(get_links())

@app.route("/api/links/<int:link_id>", methods=["DELETE"])
def remove_link(link_id):
    delete_link(link_id)
    return ok()

# ── Plugin management extras ──────────────────────────────────────────────────
@app.route("/api/plugins/<plugin_id>/disable", methods=["POST"])
@require_loader
def disable_plugin(plugin_id):
    lp = _loader.get_plugin(plugin_id)
    if not lp:
        return err("Plugin not found", 404)
    lp.enabled = False
    from config import save_plugin_state
    save_plugin_state(plugin_id, enabled=False)
    return ok({"enabled": False})

@app.route("/api/plugins/<plugin_id>", methods=["DELETE"])
@require_loader
def delete_plugin(plugin_id):
    # Plugin IDs are always of the form plugin_<snake_case>. Reject anything
    # else up front so a crafted id like "../../main" or one containing a
    # path separator can't slip through into the file path.
    if not re.match(r"^[a-z0-9_]+$", plugin_id):
        return err("Invalid plugin ID format", 400)

    lp = _loader.get_plugin(plugin_id)
    if not lp:
        return err("Plugin not found", 404)

    from plugin_loader import PLUGINS_DIR
    plugins_root = Path(PLUGINS_DIR).resolve()
    plugin_file = (plugins_root / f"{plugin_id}.py").resolve()

    # Belt-and-braces: even if the regex is satisfied, verify the resolved
    # path still lives inside plugins/ before we unlink it.
    try:
        plugin_file.relative_to(plugins_root)
    except ValueError:
        return err("Invalid plugin ID", 400)

    try:
        if plugin_file.exists():
            plugin_file.unlink()
        _loader.reload_plugins()
    except Exception as e:
        return err(str(e))
    return ok({"deleted": plugin_id})

# ── Lessons ───────────────────────────────────────────────────────────────────
@app.route("/api/lessons")
def list_lessons():
    return ok(get_active_lessons())

@app.route("/api/lessons", methods=["POST"])
def create_lesson():
    data = request.get_json(force=True)
    add_lesson(data.get("lesson", ""), data.get("source", ""))
    return ok(get_active_lessons())

@app.route("/api/lessons/<int:lesson_id>", methods=["DELETE"])
def remove_lesson(lesson_id):
    delete_lesson(lesson_id)
    return ok()

@app.route("/api/lessons/<int:lesson_id>/toggle", methods=["POST"])
def toggle_lesson_route(lesson_id):
    data = request.get_json(force=True)
    toggle_lesson(lesson_id, data.get("active", True))
    return ok()

# ── Style preferences ─────────────────────────────────────────────────────────
@app.route("/api/style")
def get_style():
    return ok({"content": get_style_preferences()})

@app.route("/api/style", methods=["POST"])
def save_style():
    data = request.get_json(force=True)
    save_style_preferences(data.get("content", ""))
    return ok()

# ── Knowledge Base ────────────────────────────────────────────────────────────
@app.route("/api/knowledge")
def list_knowledge():
    return ok(get_knowledge_entries())


@app.route("/api/knowledge", methods=["POST"])
def create_knowledge():
    data = request.get_json(force=True) or {}
    category = (data.get("category") or "").strip()
    title    = (data.get("title") or "").strip()
    content  = data.get("content") or ""
    if not category or not title:
        return err("'category' and 'title' are required")
    enabled = int(data.get("enabled", 1))
    entry_id = add_knowledge_entry(category, title, content, enabled)
    return ok({"id": entry_id})


@app.route("/api/knowledge/<int:entry_id>", methods=["PUT"])
def update_knowledge(entry_id):
    data = request.get_json(force=True) or {}
    update_knowledge_entry(
        entry_id,
        category=data.get("category"),
        title=data.get("title"),
        content=data.get("content"),
        enabled=data.get("enabled"),
    )
    return ok({"id": entry_id, "updated": True})


@app.route("/api/knowledge/<int:entry_id>", methods=["DELETE"])
def remove_knowledge(entry_id):
    delete_knowledge_entry(entry_id)
    return ok({"id": entry_id, "deleted": True})


# ── BAS Clients ───────────────────────────────────────────────────────────────
BAS_CSV_HEADERS = [
    "client_name", "entity_name", "abn", "frequency", "client_email",
    "last_data_received", "last_reminder_sent", "status", "notes",
]


def _bas_csv_template() -> str:
    return ",".join(BAS_CSV_HEADERS) + "\n"


@app.route("/api/bas-clients")
def list_bas_clients():
    return ok(get_bas_clients())


@app.route("/api/bas-clients", methods=["POST"])
def create_bas_client():
    data = request.get_json(force=True) or {}
    if not (data.get("client_name") or "").strip():
        return err("'client_name' is required")
    try:
        new_id = add_bas_client(data)
    except ValueError as e:
        return err(str(e))
    return ok(get_bas_client(new_id))


@app.route("/api/bas-clients/<int:client_id>", methods=["PUT"])
def update_bas_client_endpoint(client_id):
    data = request.get_json(force=True) or {}
    update_bas_client(client_id, data)
    updated = get_bas_client(client_id)
    if not updated:
        return err("Not found", 404)
    return ok(updated)


@app.route("/api/bas-clients/<int:client_id>", methods=["DELETE"])
def delete_bas_client_endpoint(client_id):
    delete_bas_client(client_id)
    return ok({"id": client_id, "deleted": True})


@app.route("/api/bas-clients/upload", methods=["POST"])
def upload_bas_clients():
    """Parse a CSV upload and replace the bas_clients table with its contents."""
    import csv
    import io

    if "file" not in request.files:
        return err("No file provided")
    file = request.files["file"]
    if not file.filename:
        return err("No filename")
    if not file.filename.lower().endswith(".csv"):
        return err("File must be .csv")

    try:
        raw = file.read()
        if len(raw) > MAX_UPLOAD_SIZE:
            return err("File too large", 413)
        text = raw.decode("utf-8-sig", errors="replace")
        reader = csv.DictReader(io.StringIO(text))
        if not reader.fieldnames:
            return err("CSV has no header row")
        # Normalise header names to lower_snake
        fieldmap = {name: name.strip().lower().replace(" ", "_") for name in reader.fieldnames}
        rows: list[dict] = []
        for raw_row in reader:
            mapped = {fieldmap[k]: (v or "").strip() for k, v in raw_row.items() if k}
            if not mapped.get("client_name"):
                continue
            rows.append({k: mapped.get(k) for k in BAS_CLIENT_COLUMNS})
        if not rows:
            return err("No valid rows found (client_name is required)")
        count = bulk_replace_bas_clients(rows)
    except Exception as e:
        return err(f"Failed to parse CSV: {e}")
    return ok({"imported": count, "clients": get_bas_clients()})


@app.route("/api/bas-clients/template", methods=["GET"])
def bas_clients_template():
    """Return a blank CSV with the correct headers."""
    return Response(
        _bas_csv_template(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=bas_clients_template.csv"},
    )


@app.route("/api/bas-dates", methods=["GET"])
def get_bas_dates_endpoint():
    """Return calculated BAS dates for current and next financial year."""
    from bas_dates import (
        get_bas_dates,
        get_next_due_quarter,
        get_upcoming_deadlines,
        get_financial_year,
    )

    fy = get_financial_year()
    current_dates = get_bas_dates(fy)
    next_dates = get_bas_dates(fy + 1)
    next_due = get_next_due_quarter()
    upcoming = get_upcoming_deadlines(days_ahead=60)

    def serialise(q):
        return {
            "quarter": q["quarter"],
            "period": q["period"],
            "period_start": q["period_start"].isoformat(),
            "period_end": q["period_end"].isoformat(),
            "standard_due": q["standard_due"].isoformat(),
            "agent_due": q["agent_due"].isoformat(),
            "has_extension": q["has_extension"],
            "data_request_by": q["data_request_by"].isoformat(),
            "description": q["description"],
        }

    return jsonify({
        "ok": True,
        "financial_year": f"FY{fy-1}-{str(fy)[2:]}",
        "current_year": [serialise(q) for q in current_dates],
        "next_year": [serialise(q) for q in next_dates],
        "next_due": serialise(next_due) if next_due else None,
        "upcoming_60_days": [serialise(q) for q in upcoming],
    })


# ── Chat history ──────────────────────────────────────────────────────────────
@app.route("/api/chat/history")
def chat_history():
    return ok(get_feedback_history(200))

@app.route("/api/chat/history", methods=["DELETE"])
def clear_chat_history():
    clear_feedback_history()
    return ok()


# ── Chat export (Word .docx) ──────────────────────────────────────────────────

def _add_formatted_text(paragraph, text: str):
    """Parse inline markdown formatting (**bold**, *italic*) into Word runs."""
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _markdown_to_docx(
    markdown_text: str,
    title: str,
    client_name: str | None = None,
    entity_name: str | None = None,
):
    """Convert markdown-formatted text into a styled python-docx Document.

    Returns the Document object so callers can extend it before saving.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)

    if client_name:
        p = doc.add_paragraph()
        p.add_run(f"Client: {client_name}").bold = True
        if entity_name:
            p.add_run(f"  —  {entity_name}")

    doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph("")  # spacer

    lines = (markdown_text or "").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith("|") and "|" in stripped[1:]:
            # Collect contiguous markdown table rows.
            table_lines = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            # Drop the separator row (--- | ---) if present.
            rows = [
                [c.strip() for c in r.strip("|").split("|")]
                for r in table_lines
                if not re.match(r"^\|?\s*[-:|\s]+\|?\s*$", r)
            ]
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Light List"
                for r_idx, row in enumerate(rows):
                    cells = table.rows[r_idx].cells
                    for c_idx in range(cols):
                        text = row[c_idx] if c_idx < len(row) else ""
                        cell_para = cells[c_idx].paragraphs[0]
                        if r_idx == 0:
                            run = cell_para.add_run(text)
                            run.bold = True
                        else:
                            _add_formatted_text(cell_para, text)
            continue  # already advanced i
        elif stripped == "---":
            doc.add_paragraph("_" * 50)
        elif stripped:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)
        else:
            doc.add_paragraph("")
        i += 1

    for section in doc.sections:
        footer_para = section.footer.paragraphs[0]
        footer_para.text = "MC & S Pty Ltd — Confidential"
        for run in footer_para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)

    return doc


def _docx_response(doc, filename: str):
    """Serialise a python-docx Document and return a Flask file response."""
    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


def _last_assistant_message(messages: list) -> str:
    for m in reversed(messages or []):
        if isinstance(m, dict) and m.get("role") == "assistant":
            content = m.get("content", "")
            return content if isinstance(content, str) else str(content)
    return ""


def _has_structured_recommendation(text: str) -> bool:
    if not text:
        return False
    return bool(re.search(r"^#{1,3}\s", text, re.MULTILINE))


@app.route("/api/chat/export", methods=["POST"])
def chat_export():
    """Render the supplied conversation as a Word document.

    Body:
      agent_name:   str
      messages:     [{"role", "content"}, ...]
      client_name:  optional
      entity_name:  optional
      export_type:  "transcript" (default) | "summary" | "recommendation"
    """
    from docx import Document
    from docx.shared import Pt

    body = request.get_json(silent=True) or {}
    agent_name = (body.get("agent_name") or "Assistant").strip() or "Assistant"
    messages = body.get("messages") or []
    client_name = (body.get("client_name") or "").strip() or None
    entity_name = (body.get("entity_name") or "").strip() or None
    export_type = (body.get("export_type") or "transcript").lower()
    if export_type not in ("transcript", "summary", "recommendation"):
        return err("Invalid export_type — use transcript, summary, or recommendation")
    if not isinstance(messages, list) or not messages:
        return err("No messages to export")

    if client_name:
        try:
            from client_utils import normalise_client_name
            client_name = normalise_client_name(client_name)
        except Exception:
            pass

    safe_agent = re.sub(r"[^a-z0-9]+", "_", agent_name.lower()).strip("_") or "chat"
    ts = datetime.now().strftime("%Y-%m-%d_%H%M")

    if export_type == "summary":
        try:
            import anthropic as anthropic_lib
            api_key = get_setting("anthropic_api_key", "")
            if not api_key:
                return err("Anthropic API key not configured")
            client = anthropic_lib.Anthropic(api_key=api_key)
            convo_text = "\n\n".join(
                f"{m.get('role', 'user').upper()}: {m.get('content', '')}"
                for m in messages
                if isinstance(m, dict)
            )
            summary_resp = client.messages.create(
                model=get_claude_model_reasoning(),
                max_tokens=2048,
                system=(
                    "Summarise the following conversation into a concise professional "
                    "document suitable for filing in a client folder. Include key "
                    "decisions, advice given, and action items. Format with clear "
                    "markdown headings (##, ###), bullet points, and bold for emphasis."
                ),
                messages=[{"role": "user", "content": convo_text[:80000]}],
            )
            summary_md = summary_resp.content[0].text
        except Exception as e:
            return err(f"Summary generation failed: {e}")

        doc = _markdown_to_docx(
            summary_md,
            title=f"{agent_name} — Conversation Summary",
            client_name=client_name,
            entity_name=entity_name,
        )
        return _docx_response(doc, f"{safe_agent}_summary_{ts}.docx")

    if export_type == "recommendation":
        last = _last_assistant_message(messages)
        if not _has_structured_recommendation(last):
            return err(
                "No structured recommendation to export — ask the specialist to "
                "produce one first.",
                400,
            )
        doc = _markdown_to_docx(
            last,
            title=f"{agent_name} — Recommendation",
            client_name=client_name,
            entity_name=entity_name,
        )
        return _docx_response(doc, f"{safe_agent}_recommendation_{ts}.docx")

    # Default: transcript
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    doc.add_heading(f"{agent_name} — Chat Export", level=1)
    date_str = datetime.now().strftime("%d %B %Y, %I:%M %p")
    subtitle = doc.add_paragraph()
    subtitle.add_run(f"Date: {date_str}  |  Agent: {agent_name}").italic = True

    if client_name:
        client_para = doc.add_paragraph()
        run = client_para.add_run(f"Client: {client_name}")
        run.bold = True
        if entity_name:
            client_para.add_run(f"  |  Entity: {entity_name}")

    doc.add_paragraph()  # blank spacer

    for m in messages:
        if not isinstance(m, dict):
            continue
        role = m.get("role")
        content = m.get("content", "")
        if not isinstance(content, str):
            content = str(content)
        label = "User:" if role == "user" else f"{agent_name}:"
        para = doc.add_paragraph()
        label_run = para.add_run(label + " ")
        label_run.bold = True
        para.add_run(content)
        doc.add_paragraph()  # blank line between turns

    for section in doc.sections:
        footer_para = section.footer.paragraphs[0]
        footer_para.text = "MC & S Pty Ltd — Confidential"
        for run in footer_para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)

    return _docx_response(doc, f"{safe_agent}_{ts}.docx")


# ── Microsoft OAuth callback ─────────────────────────────────────────────────
@app.route("/auth/callback")
def auth_callback():
    """Receives the Microsoft OAuth2 redirect and passes the code to GraphClient."""
    code = request.args.get("code")
    error = request.args.get("error")
    if _graph_client is not None:
        _graph_client.receive_auth_code(code=code, error=error)
    if error:
        safe_error = html.escape(error)
        body = f"""<html><body style='font-family:Arial;text-align:center;padding:60px'>
<h2 style='color:#c0392b'>Authentication Failed</h2>
<p>{safe_error}</p>
<p>You can close this tab.</p>
</body></html>"""
        return body, 400
    body = """<html><body style='font-family:Arial;text-align:center;padding:60px'>
<h2 style='color:#2E7D32'>&#10003; Authentication Successful</h2>
<p>You can close this tab and return to MC&amp;S CoWorker.</p>
<script>setTimeout(function(){window.close();},2000);</script>
</body></html>"""
    return body, 200


# ── Frontend static file serving ──────────────────────────────────────────────────────────────
# Serves the built React app (frontend_dist/) when running in installed mode.

def _frontend_dir() -> str | None:
    candidate = os.path.join(os.path.dirname(os.path.abspath(__file__)), "frontend_dist")
    return candidate if os.path.isdir(candidate) else None


@app.route("/", defaults={"path": ""})
@app.route("/<path:path>")
def serve_frontend(path):
    """Serve the React SPA. Non-API paths return index.html (SPA routing)."""
    frontend = _frontend_dir()
    if frontend is None:
        return jsonify({"error": "Frontend not built."}), 404
    # Serve real static assets (JS, CSS, images, etc.)
    if path and os.path.isfile(os.path.join(frontend, path)):
        return send_from_directory(frontend, path)
    # All other paths → index.html (React Router handles routing client-side)
    return send_from_directory(frontend, "index.html")


def run_server(host="127.0.0.1", port=API_PORT, debug=False):
    """Start the Flask server in a background thread."""
    app.run(host=host, port=port, debug=debug, use_reloader=False, threaded=True)


def start_in_thread(host="127.0.0.1", port=API_PORT):
    # Wire SSE broadcaster to EventBus before starting
    try:
        _wire_sse_to_event_bus()
    except Exception:
        pass  # EventBus may not be ready yet — wiring happens lazily
    t = threading.Thread(target=run_server, args=(host, port), daemon=True)
    t.start()
    return t
