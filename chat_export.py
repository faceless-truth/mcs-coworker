"""Generate a Word document from a chat_messages session for SharePoint archive.

Uses python-docx, which v2.3 already depends on for conversation export.
Output is a single .docx file in memory (BytesIO) ready for upload.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Iterable

from docx import Document
from docx.shared import Pt, RGBColor


def build_chat_docx(
    *,
    specialist_display_name: str,
    title: str,
    messages: Iterable[dict],
    user_display_name: str,
    archived_at: datetime,
) -> io.BytesIO:
    """Build a .docx containing the full conversation. Return BytesIO at pos 0."""
    doc = Document()

    doc.add_heading(title, level=1)
    meta = doc.add_paragraph()
    meta.add_run(f"Specialist: {specialist_display_name}\n").bold = True
    meta.add_run(f"User: {user_display_name}\n")
    meta.add_run(f"Archived: {archived_at.strftime('%Y-%m-%d %H:%M')}")

    doc.add_paragraph()

    for msg in messages:
        role = msg["role"]
        content = msg["content"]
        created = msg.get("created_at")

        para = doc.add_paragraph()
        run = para.add_run(_role_label(role, user_display_name))
        run.bold = True
        if isinstance(created, str):
            try:
                created_dt = datetime.fromisoformat(created.replace("Z", "+00:00"))
                stamp = para.add_run(f"  ({created_dt.strftime('%H:%M')})")
                stamp.font.size = Pt(9)
                stamp.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            except ValueError:
                pass

        body = doc.add_paragraph(content)
        body.paragraph_format.left_indent = Pt(18)
        body.paragraph_format.space_after = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _role_label(role: str, user_display_name: str) -> str:
    if role == "user":
        return f"{user_display_name}:"
    if role == "assistant":
        return "MCS CoWorker:"
    return f"{role}:"
